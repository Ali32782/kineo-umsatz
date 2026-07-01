import os, zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from sqlalchemy.orm import Session
from datetime import datetime

from database import UmsatzData, MonthlyInput, MAStammdaten, User, BilatData
from calc import compute_zeg, compute_soll_tage, zeg_color, MONTH_NAMES_DE, is_employed_in_month
from schedule_utils import get_schedule_entries_for_month
from standort_calc import expand_ma_standort_rows

EXPORTS_DIR = os.environ.get("EXPORTS_DIR", os.path.join(os.path.dirname(__file__), "../exports"))
TEAL = RGBColor(0x00, 0x6B, 0x6B)

KAT_LABELS = {
    "a": "Kat. A — Profitabilität & Auslastung",
    "b": "Kat. B — Qualität & Operational Excellence",
    "c": "Kat. C — Satisfaction Intern – Team & Kultur",
    "d": "Kat. D — Satisfaction Extern – Patienten & Zuweiser",
}

def zeg_label(zeg_b):
    if zeg_b is None: return "—"
    pct = round(zeg_b * 100, 1)
    emoji = "✓" if zeg_b >= 1.0 else ("~" if zeg_b >= 0.85 else "!")
    return f"{emoji} {pct}%"

def _half_for_month(month: int) -> int:
    return 1 if month <= 6 else 2

def generate_bilats_zip(year: int, month: int, db: Session, current_user: User, period_label: str = None) -> str:
    all_mas = db.query(MAStammdaten).all()
    mas = [
        m for m in all_mas
        if any(is_employed_in_month(m.eintritt, m.austritt, year, mo, m.is_active) for mo in range(1, month + 1))
    ]
    if current_user.role == "teamlead" and current_user.team:
        mas = [m for m in mas if m.team == current_user.team]

    umsatz_all = {}
    inputs_all = {}
    for m in range(1, month + 1):
        for r in db.query(UmsatzData).filter(UmsatzData.year == year, UmsatzData.month == m).all():
            umsatz_all[(r.ma_name, m)] = r.umsatz
        for r in db.query(MonthlyInput).filter(MonthlyInput.year == year, MonthlyInput.month == m).all():
            inputs_all[(r.ma_name, m)] = r

    # Load bilat data by period_label (new) or half (legacy fallback)
    if period_label:
        bilat_all = {b.ma_name: b for b in db.query(BilatData).filter_by(year=year, period_label=period_label).all()}
    else:
        half = _half_for_month(month)
        bilat_all = {b.ma_name: b for b in db.query(BilatData).filter_by(year=year, half=half).all()}

    os.makedirs(EXPORTS_DIR, exist_ok=True)
    zip_path = os.path.join(EXPORTS_DIR, f"Bilats_{MONTH_NAMES_DE[month]}_{year}.zip")

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for ma in mas:
            bilat = bilat_all.get(ma.name)
            doc_path = _generate_bilat_doc(ma, year, month, umsatz_all, inputs_all, bilat, db=db)
            zf.write(doc_path, os.path.basename(doc_path))
            os.remove(doc_path)

    return zip_path

def _generate_bilat_doc(ma: MAStammdaten, year: int, month: int,
                         umsatz_all: dict, inputs_all: dict, bilat: "BilatData" = None, db=None) -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    title = doc.add_heading('', level=0)
    title.clear()
    run = title.add_run(f'KINEO AG  |  Bilateral HJ1 {year}')
    run.font.color.rgb = TEAL
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.add_run('Mitarbeiter/in: ').bold = True
    p.add_run(ma.display_name or ma.name)
    p = doc.add_paragraph()
    r = p.add_run('Team: ')
    r.bold = True
    p.add_run(ma.team or '—')
    p = doc.add_paragraph()
    r = p.add_run('Zeitraum: ')
    r.bold = True
    p.add_run(f'Januar {year} – {MONTH_NAMES_DE[month]} {year}')
    doc.add_paragraph('')

    # ── ZEG-B Performance Table ──────────────────────────────────────────
    h = doc.add_heading('Umsatz & Zielerreichungsgrad (ZEG-B)', level=2)
    h.runs[0].font.color.rgb = TEAL

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0].cells
    for i, txt in enumerate(['Monat', 'Soll-Tage', 'Umsatz CHF', 'Prod-Tage (B)', 'ZEG-B']):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = TEAL

    zeg_b_values = []
    for m in range(1, month + 1):
        if not is_employed_in_month(ma.eintritt, ma.austritt, year, m, ma.is_active):
            continue
        umsatz = umsatz_all.get((ma.name, m), 0)
        inp = inputs_all.get((ma.name, m))
        soll = compute_soll_tage(ma.name, year, m, db=db)
        if soll == 0 and umsatz == 0:
            continue
        zeg = compute_zeg(
            ma.name, year, m, umsatz,
            ferien_t=inp.ferien_t if inp else 0,
            kurs_h=inp.kurs_h if inp else 0,
            workshop_h=inp.workshop_h if inp else 0,
            marketing_h=inp.marketing_h if inp else 0,
            laufanalyse_h=inp.laufanalyse_h if inp else 0,
            krank_t=inp.krank_t if inp else 0,
            db=db,
        )
        row = tbl.add_row().cells
        row[0].text = MONTH_NAMES_DE[m]
        row[1].text = str(soll)
        row[2].text = f"CHF {umsatz:,.0f}".replace(',', "'")
        row[3].text = f"{zeg['prod_b']:.1f}"
        row[4].text = zeg_label(zeg['zeg_b'])
        if zeg['zeg_b']:
            zeg_b_values.append(zeg['zeg_b'])

    if zeg_b_values:
        avg = sum(zeg_b_values) / len(zeg_b_values)
        row = tbl.add_row().cells
        row[0].text = f'Ø {month} Monate'
        row[0].paragraphs[0].runs[0].bold = True
        row[4].text = zeg_label(avg)
        row[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph('')
    doc.add_paragraph('Legende: ✓ ≥ 100%  |  ~ 85–99%  |  ! < 85%  |  Ziel: CHF 1\'040 / produktiver Tag')
    doc.add_paragraph('')

    # ── Standort-Aufteilung (letzter Monat) ─────────────────────────────
    h3 = doc.add_heading(f'Standort-Aufteilung {MONTH_NAMES_DE[month]} {year}', level=2)
    h3.runs[0].font.color.rgb = TEAL
    umsatz_m = umsatz_all.get((ma.name, month), 0)
    inp_m = inputs_all.get((ma.name, month))
    if is_employed_in_month(ma.eintritt, ma.austritt, year, month, ma.is_active) and umsatz_m:
        zeg_m = compute_zeg(
            ma.name, year, month, umsatz_m,
            ferien_t=inp_m.ferien_t if inp_m else 0,
            kurs_h=inp_m.kurs_h if inp_m else 0,
            workshop_h=inp_m.workshop_h if inp_m else 0,
            marketing_h=inp_m.marketing_h if inp_m else 0,
            laufanalyse_h=inp_m.laufanalyse_h if inp_m else 0,
            krank_t=inp_m.krank_t if inp_m else 0,
            db=db,
        )
        sched = get_schedule_entries_for_month(ma.name, year, month, db)
        row_data = {"name": ma.name, "display_name": ma.display_name, "team": ma.team, "umsatz": umsatz_m, **zeg_m}
        standort_rows = expand_ma_standort_rows(row_data, ma.bg_pct, ma.team, sched)
        tbl3 = doc.add_table(rows=1, cols=4)
        tbl3.style = 'Table Grid'
        for i, txt in enumerate(['Standort', 'FTE-Anteil', 'Umsatz CHF', 'ZEG-B']):
            tbl3.rows[0].cells[i].text = txt
            tbl3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for sr in standort_rows:
            row = tbl3.add_row().cells
            row[0].text = sr['team'] + (' (kein Umsatz)' if sr.get('is_office') else '')
            row[1].text = f"{(sr['bg_pct']*100):.0f}%"
            row[2].text = f"CHF {sr['umsatz']:,.0f}".replace(',', "'") if sr['umsatz'] else '—'
            row[3].text = zeg_label(sr.get('zeg_b'))
    else:
        doc.add_paragraph('Kein Umsatz im Berichtsmonat.').runs[0].italic = True

    doc.add_paragraph('')

    # ── Erfasste Bewertung (aus BilatData, falls vorhanden) ──────────────
    h2 = doc.add_heading('Bewertung nach Kategorie', level=2)
    h2.runs[0].font.color.rgb = TEAL

    if bilat is None:
        p = doc.add_paragraph('Noch keine Bewertung erfasst. Bitte im Bereich "Bilaterals" in der App ausfüllen.')
        p.runs[0].italic = True
    else:
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = 'Table Grid'
        hdr2 = tbl2.rows[0].cells
        for i, txt in enumerate(['Kategorie', 'Selbst (1–5)', 'FK (1–5)', 'Kommentar']):
            hdr2[i].text = txt
            hdr2[i].paragraphs[0].runs[0].bold = True
            hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)
            hdr2[i].paragraphs[0].runs[0].font.color.rgb = TEAL

        for k in ["a", "b", "c", "d"]:
            self_v = getattr(bilat, f"kat_{k}_self", None)
            fk_v = getattr(bilat, f"kat_{k}_fk", None)
            comment = getattr(bilat, f"kat_{k}_comment", None) or ''
            row2 = tbl2.add_row().cells
            row2[0].text = KAT_LABELS[k]
            row2[1].text = str(self_v) if self_v is not None else '—'
            row2[2].text = str(fk_v) if fk_v is not None else '—'
            row2[3].text = comment

        doc.add_paragraph('')

        # Themen des Mitarbeiters
        h3 = doc.add_heading('Themen des Mitarbeiters', level=3)
        h3.runs[0].font.color.rgb = TEAL
        doc.add_paragraph(bilat.themen_ma or '—')

        # Vereinbarungen & nächste Schritte
        h4 = doc.add_heading('Vereinbarungen & nächste Schritte', level=3)
        h4.runs[0].font.color.rgb = TEAL
        doc.add_paragraph(bilat.vereinbarungen or '—')

        # Abschluss
        h5 = doc.add_heading('Abschluss', level=3)
        h5.runs[0].font.color.rgb = TEAL
        p = doc.add_paragraph()
        p.add_run('Gesprächseindruck: ').bold = True
        p.add_run(bilat.gespraechseindruck or '—')
        p = doc.add_paragraph()
        p.add_run('Nächstes Bilat-Datum: ').bold = True
        p.add_run(bilat.naechstes_bilat or '—')

    os.makedirs(EXPORTS_DIR, exist_ok=True)
    safe_name = ma.name.replace('.', '_').replace(' ', '_')
    path = os.path.join(EXPORTS_DIR, f"Bilat_{safe_name}_{MONTH_NAMES_DE[month]}_{year}.docx")
    doc.save(path)
    return path


def generate_single_bilat(year: int, month: int, ma_name: str, db, period_label: str = None) -> str:
    """Generate a single bilat Word doc for one MA, including erfasste Bewertungen."""
    ma = db.query(MAStammdaten).filter_by(name=ma_name).first()
    if not ma:
        raise ValueError(f"MA {ma_name} not found")
    if not is_employed_in_month(ma.eintritt, ma.austritt, year, month, ma.is_active):
        raise ValueError(f"MA {ma_name} war im Berichtsmonat nicht angestellt")
    umsatz_all = {(r.ma_name, r.month): r.umsatz
                  for r in db.query(UmsatzData).filter(UmsatzData.year==year).all()}
    inputs_all = {(r.ma_name, r.month): r
                  for r in db.query(MonthlyInput).filter(MonthlyInput.year==year).all()}
    # Load by period_label (new) or half (legacy)
    if period_label:
        bilat = db.query(BilatData).filter_by(ma_name=ma_name, year=year, period_label=period_label).first()
    else:
        half = _half_for_month(month)
        bilat = db.query(BilatData).filter_by(ma_name=ma_name, year=year, half=half).first()
    return _generate_bilat_doc(ma, year, month, umsatz_all, inputs_all, bilat, db=db)
