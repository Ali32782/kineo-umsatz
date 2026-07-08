"""Bilateral Word export — pro MA individuelle HJ1-Vorlage mit vollen Rubriken."""
from __future__ import annotations

import os
import shutil
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from sqlalchemy.orm import Session

from bilat_template_map import resolve_bilat_template
from calc import MONTH_NAMES_DE, ZIEL, compute_soll_tage, compute_zeg, is_employed_in_month, zeg_color
from database import BilatData, MAStammdaten, User

ROLE_LABELS_DE = {
    "therapeut": "Therapeut/in",
    "sl": "Standortlead",
    "teamlead": "Teamlead",
    "bd": "Business Development",
    "management": "Management",
    "ceo": "CEO",
    "coo": "COO",
}

MONTH_SHORT = {1: "Jan", 2: "Feb", 3: "Mrz", 4: "Apr", 5: "Mai", 6: "Jun",
               7: "Jul", 8: "Aug", 9: "Sep", 10: "Okt", 11: "Nov", 12: "Dez"}


def _half_label(month: int, year: int) -> str:
    return f"{'1.' if month <= 6 else '2.'} HJ {year}"


def _period_range(month: int, year: int) -> str:
    return f"Jan – {MONTH_NAMES_DE[month]} {year}"


def _fmt_chf(amount: float) -> str:
    return f"CHF {round(amount):,}".replace(",", "'")


def _zeg_pct(zeg_b: float | None) -> str:
    if zeg_b is None:
        return "—"
    return f"{round(zeg_b * 100, 1)}%"


def _vs_ziel(zeg_b: float | None) -> str:
    if zeg_b is None:
        return "—"
    diff = round(zeg_b * 100 - 100, 1)
    sign = "+" if diff >= 0 else ""
    return f"{sign}{diff} Pkt."


def _dash_val(value: float) -> str:
    if not value:
        return "—"
    return f"-{value:g}"


_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_TEAL = RGBColor(0x00, 0xB4, 0xA6)
_RATING_GRAY = RGBColor(0x44, 0x44, 0x44)

_ZEG_FILL = {"green": "D5F5E3", "amber": "FDEBD0", "red": "FADBD8"}
_HEADER_FILL = "2C3E50"
_LABEL_FILL = "F5F5F5"


def _zeg_fill_hex(zeg_b: float | None) -> str | None:
    return _ZEG_FILL.get(zeg_color(zeg_b))


def _clone_tc_pr(source, target) -> None:
    src_pr = source._tc.tcPr
    if src_pr is None:
        return
    dst_pr = target._tc.tcPr
    if dst_pr is not None:
        target._tc.remove(dst_pr)
    target._tc.insert(0, deepcopy(src_pr))


def _set_cell_shading(cell, fill_hex: str | None) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if not fill_hex:
        if shd is not None:
            tc_pr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)


def _set_row_cant_split(row, *, keep_with_next: bool = False) -> None:
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
    if keep_with_next:
        for para in row.cells[0].paragraphs:
            p_pr = para._p.get_or_add_pPr()
            if p_pr.find(qn("w:keepNext")) is None:
                p_pr.append(OxmlElement("w:keepNext"))


W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"


def _set_rating_cell(cell, value: int | None) -> None:
    """Word-Checkboxen (w:sdt) der Vorlage markieren und Zahl hervorheben."""
    if len(cell.paragraphs) < 2:
        return
    para = cell.paragraphs[1]
    children = list(para._p)
    idx = 0
    while idx < len(children):
        child = children[idx]
        if child.tag != qn("w:sdt"):
            idx += 1
            continue
        checkbox = child.find(f".//{{{W14_NS}}}checkbox")
        if checkbox is None:
            idx += 1
            continue

        num = None
        num_run_el = None
        j = idx + 1
        while j < len(children):
            sibling = children[j]
            if sibling.tag == qn("w:r"):
                text_el = sibling.find(qn("w:t"))
                if text_el is not None and text_el.text and text_el.text.strip().isdigit():
                    num = int(text_el.text.strip())
                    num_run_el = sibling
                    break
            elif sibling.tag == qn("w:sdt"):
                break
            j += 1

        if num is not None:
            selected = value is not None and num == value
            checked = checkbox.find(f"{{{W14_NS}}}checked")
            if checked is not None:
                checked.set(f"{{{W14_NS}}}val", "1" if selected else "0")
            for text_el in child.iter(qn("w:t")):
                if text_el.text in ("☐", "☑"):
                    text_el.text = "☑" if selected else "☐"
            if num_run_el is not None:
                from docx.text.run import Run

                run = Run(num_run_el, para)
                run.font.bold = selected
                run.font.color.rgb = _TEAL if selected else _RATING_GRAY
            idx = j + 1
        else:
            idx += 1


def _save_row_tc_pr(table, row_idx: int) -> list:
    if row_idx >= len(table.rows):
        return []
    return [deepcopy(c._tc.tcPr) if c._tc.tcPr is not None else None for c in table.rows[row_idx].cells]


def _apply_saved_tc_pr(cell, saved: list, idx: int, fallback: int = 0) -> None:
    src_idx = min(idx, len(saved) - 1) if saved else fallback
    if saved and 0 <= src_idx < len(saved) and saved[src_idx] is not None:
        dst_pr = cell._tc.tcPr
        if dst_pr is not None:
            cell._tc.remove(dst_pr)
        cell._tc.insert(0, deepcopy(saved[src_idx]))


def _cell_fill_hex(cell) -> str | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _is_dark_fill(fill: str | None) -> bool:
    if not fill or fill.lower() in ("auto", "ffffff"):
        return False
    if len(fill) == 6:
        try:
            r, g, b = int(fill[0:2], 16), int(fill[2:4], 16), int(fill[4:6], 16)
            return (r * 0.299 + g * 0.587 + b * 0.114) < 140
        except ValueError:
            return False
    return True


def _set_cell_text(cell, text: str, *, section_header: bool = False) -> None:
    """Text setzen ohne Banner-Formatierung (weiss auf dunkel) zu zerstören."""
    value = str(text) if text is not None else ""
    dark = section_header or _is_dark_fill(_cell_fill_hex(cell))

    if dark:
        cell.text = ""
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = para.add_run(value.strip())
        run.font.color.rgb = _WHITE
        run.font.bold = True
        run.font.size = Pt(11)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        return

    if cell.paragraphs and cell.paragraphs[0].runs:
        para = cell.paragraphs[0]
        para.runs[0].text = value
        for extra in para.runs[1:]:
            extra.text = ""
    else:
        cell.text = value


def _merge_row_cells(row, start: int = 0, end: int | None = None) -> None:
    end = end if end is not None else len(row.cells) - 1
    if start < end:
        row.cells[start].merge(row.cells[end])


def _delete_row(table, row_idx: int) -> None:
    table._tbl.remove(table.rows[row_idx]._tr)


def _clear_rows(table, keep: int = 1) -> None:
    while len(table.rows) > keep:
        _delete_row(table, len(table.rows) - 1)


def _row_text(table, row_idx: int = 0) -> str:
    if row_idx >= len(table.rows):
        return ""
    return " ".join(c.text.strip() for c in table.rows[row_idx].cells)


def _find_tables(doc: Document) -> dict:
    found: dict = {}
    for table in doc.tables:
        r0 = _row_text(table, 0)
        r1 = _row_text(table, 1) if len(table.rows) > 1 else ""
        if "Mitarbeiter/in" in r0 and "Standort" in r0:
            found["header"] = table
        elif r0.strip().startswith("PERFORMANCE") and "ZEG-B" in r0:
            found["perf_title"] = table
        elif r0.startswith("ZEG-B % des Ziels") or r1.startswith("Auslastung vs. Ziel"):
            found["zeg_matrix"] = table
        elif r0.startswith("QUALITATIVE ZIELE"):
            found["qual_title"] = table
        elif r0.startswith("Ziel 1. HJ") or r0.startswith("Ziel 2. HJ"):
            found["qual_goals"] = table
        elif "BEWERTUNG NACH KATEGORIE" in r0:
            found["ratings_title"] = table
        elif r0.startswith("Kat.") and "Bereich" in r0:
            found["ratings"] = table
        elif "THEMEN DES MITARBEITERS" in r0:
            found["themen_title"] = table
        elif "Was liegt mir auf dem Herzen" in r0:
            found["themen"] = table
        elif "VEREINBARUNGEN" in r0 and "NÄCHSTE" in r0:
            found["vereinbarungen_title"] = table
        elif r0.startswith("Massnahme / Ziel"):
            found["vereinbarungen"] = table
        elif r0.startswith("ABSCHLUSS"):
            found["abschluss_title"] = table
        elif "Nächstes Bilat-Datum" in r0:
            found["abschluss"] = table
        elif r0.startswith("FAKTENBLATT"):
            found["faktenblatt"] = table
        elif "PERFORMANCE-DETAIL" in r0:
            found["perf_detail_title"] = table
        elif r0.startswith("Monat") and "ZEG-B %" in r0 and len(table.rows[0].cells) <= 6:
            found["perf_detail"] = table
        elif "BERECHNUNGSGRUNDLAGE" in r0:
            found["calc_title"] = table
        elif r0.startswith("Monat") and len(table.rows[0].cells) >= 10:
            found["calc_detail"] = table
        elif "GESPRÄCHSLEITFADEN" in r0:
            found["leitfaden_title"] = table
        elif r0.startswith("Performance-Einschätzung"):
            found["leitfaden"] = table
    return found


def _find_fk(db: Session, team: str | None, fk_username: str | None = None) -> User | None:
    if fk_username:
        fk = db.query(User).filter_by(username=fk_username).first()
        if fk:
            return fk
    if not team or team in ("Management", "Office"):
        return db.query(User).filter(User.role.in_(("coo", "ceo"))).first()
    lead = db.query(User).filter(User.role == "teamlead", User.team == team).first()
    if lead:
        return lead
    return db.query(User).filter(User.role == "sl", User.team == team).first()


def _trend_arrow(values: list[float]) -> str:
    if len(values) < 2:
        return "►"
    return "▲" if values[-1] > values[-2] else ("▼" if values[-1] < values[-2] else "►")


def _performance_comment(avg: float | None, perf_range: str) -> str:
    if avg is None:
        return "Zu wenig Datenpunkte für Jahresbewertung."
    pct = round(avg * 100, 1)
    if pct >= 100:
        qual = "starker 1. HJ"
    elif pct >= 90:
        qual = "solider 1. HJ – Optimierungspotenzial vorhanden"
    else:
        qual = "Optimierung im 1. HJ erforderlich"
    return f"{pct}% Ø ({perf_range}) – {qual}"


def _compute_month_row(ma, year, m, umsatz_all, inputs_all, db) -> dict | None:
    if not is_employed_in_month(ma.eintritt, ma.austritt, year, m, ma.is_active):
        return None
    umsatz = umsatz_all.get((ma.name, m), 0) or 0
    inp = inputs_all.get((ma.name, m))
    soll = compute_soll_tage(ma.name, year, m, db=db)
    if soll == 0 and umsatz == 0:
        return None
    zeg = compute_zeg(
        ma.name, year, m, umsatz,
        ferien_t=inp.ferien_t if inp else 0,
        kurs_h=inp.kurs_h if inp else 0,
        workshop_h=inp.workshop_h if inp else 0,
        marketing_h=inp.marketing_h if inp else 0,
        laufanalyse_h=inp.laufanalyse_h if inp else 0,
        krank_t=inp.krank_t if inp else 0,
        db=db,
    )
    kurs_h = (inp.kurs_h or 0) + (inp.workshop_h or 0) if inp else 0
    return {
        "month": m,
        "umsatz": umsatz,
        "soll": soll,
        "zeg": zeg,
        "ferien": inp.ferien_t if inp else 0,
        "kurs_h": kurs_h,
        "marketing": inp.marketing_h if inp else 0,
        "lauf": inp.laufanalyse_h if inp else 0,
        "krank": inp.krank_t if inp else 0,
    }


def _collect_performance(ma, year, through_month, umsatz_all, inputs_all, db) -> list[dict]:
    return [
        row for m in range(1, through_month + 1)
        if (row := _compute_month_row(ma, year, m, umsatz_all, inputs_all, db))
    ]


def _standort_funktion(ma: MAStammdaten) -> str:
    role = ROLE_LABELS_DE.get(ma.role or "", ma.role or "—")
    return f"{ma.team or '—'} / {role}"


def _reset_qual_goal_values(table) -> list[str]:
    """Zielnamen aus Vorlage behalten; Ergebnis/Status nicht in App → leeren."""
    names: list[str] = []
    for row in table.rows[1:]:
        cells = row.cells
        name = cells[0].text.strip()
        if not name or name.startswith("Ziel"):
            continue
        names.append(name)
        if len(cells) > 1:
            _set_cell_text(cells[1], "—")
        if len(cells) > 2:
            _set_cell_text(cells[2], "offen")
        if len(cells) > 3:
            _set_cell_text(cells[3], "—")
    return names


def _build_leitfaden_points(perf_range: str, qual_goal_names: list[str], bilat: BilatData | None) -> list[str]:
    points = [f"1.  Performance {perf_range}: Entwicklung & Trend besprechen"]
    for name in qual_goal_names:
        if "nicht" in name.lower() and "erfasst" in name.lower():
            continue
        points.append(f"{len(points) + 1}.  {name}")
    if bilat:
        for key, label in [("a", "Kat. A"), ("b", "Kat. B"), ("c", "Kat. C"), ("d", "Kat. D")]:
            fk_v = getattr(bilat, f"kat_{key}_fk", None)
            if fk_v is not None:
                points.append(f"{len(points) + 1}.  {label}: FK-Bewertung {fk_v}/5")
    return points


def _add_cell(row) -> None:
    from docx.oxml import OxmlElement
    row._tr.append(OxmlElement("w:tc"))


def _ensure_row_width(row, n_cols: int) -> None:
    while len(row.cells) < n_cols:
        _add_cell(row)


def _fill_zeg_matrix(table, perf_rows: list[dict], avg_zeg: float | None, bg_pct: str) -> None:
    month_cols = [r["month"] for r in perf_rows]
    zeg_values = [r["zeg"]["zeg_b"] for r in perf_rows if r["zeg"].get("zeg_b") is not None]
    headers = ["ZEG-B % des Ziels"] + [MONTH_SHORT[m] for m in month_cols] + ["Ø HJ", "Trend", "BG"]
    month_zeg = {
        m: next(r["zeg"]["zeg_b"] for r in perf_rows if r["month"] == m)
        for m in month_cols
    }
    values = ["Auslastung vs. Ziel"] + [
        _zeg_pct(month_zeg.get(m)) for m in month_cols
    ] + [_zeg_pct(avg_zeg), _trend_arrow(zeg_values), bg_pct]
    fills: list[str | None] = [None] + [_zeg_fill_hex(month_zeg.get(m)) for m in month_cols] + [
        _zeg_fill_hex(avg_zeg), None, None
    ]

    tpl_hdr = _save_row_tc_pr(table, 0)
    tpl_val = _save_row_tc_pr(table, 1)
    _clear_rows(table, keep=0)
    table.add_row()
    table.add_row()
    n_cols = len(headers)
    for row in table.rows:
        _ensure_row_width(row, n_cols)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if i == 0:
            src = 0
        elif i >= n_cols - 3:
            src = min(len(tpl_hdr) - 1, 6 + (i - (n_cols - 3)))
        else:
            src = 1
        _apply_saved_tc_pr(cell, tpl_hdr, src, fallback=0)
        _set_cell_shading(cell, _HEADER_FILL)
        _set_cell_text(cell, header, section_header=True)

    for i, (value, fill) in enumerate(zip(values, fills)):
        cell = table.rows[1].cells[i]
        if i == 0:
            src = 0
        elif i >= n_cols - 3:
            src = min(len(tpl_val) - 1, 6 + (i - (n_cols - 3)))
        else:
            src = 1
        _apply_saved_tc_pr(cell, tpl_val, src, fallback=0)
        if i == 0 or i >= n_cols - 3:
            _set_cell_shading(cell, _LABEL_FILL)
        elif fill:
            _set_cell_shading(cell, fill)
        _set_cell_text(cell, value)


def _fill_perf_detail(table, perf_rows: list[dict], year: int, perf_range: str, avg_zeg: float | None) -> None:
    tpl_rows = [_save_row_tc_pr(table, i) for i in range(1, min(3, len(table.rows)))]
    _clear_rows(table, keep=1)
    zeg_values = []
    for row_idx, pr in enumerate(perf_rows):
        zeg_b = pr["zeg"].get("zeg_b")
        if zeg_b is not None:
            zeg_values.append(zeg_b)
        row = table.add_row()
        _ensure_row_width(row, 5)
        tpl = tpl_rows[row_idx % len(tpl_rows)] if tpl_rows else []
        cells = row.cells
        row_data = [
            f"{MONTH_NAMES_DE[pr['month']]} {year}",
            _zeg_pct(zeg_b),
            _vs_ziel(zeg_b),
            f"{zeg_b:.3f}" if zeg_b is not None else "—",
            "",
        ]
        for ci, txt in enumerate(row_data):
            _apply_saved_tc_pr(cells[ci], tpl, ci, fallback=0)
            if ci == 1 and zeg_b is not None:
                fill = _zeg_fill_hex(zeg_b)
                if fill:
                    _set_cell_shading(cells[ci], fill)
            _set_cell_text(cells[ci], txt)
    if zeg_values:
        row = table.add_row()
        _ensure_row_width(row, 5)
        tpl = tpl_rows[len(perf_rows) % len(tpl_rows)] if tpl_rows else []
        avg_row = [
            f"Ø {perf_range}",
            _zeg_pct(avg_zeg),
            _vs_ziel(avg_zeg),
            f"{avg_zeg:.3f}" if avg_zeg is not None else "—",
            "",
        ]
        for ci, txt in enumerate(avg_row):
            _apply_saved_tc_pr(row.cells[ci], tpl, ci, fallback=0)
            if ci == 1 and avg_zeg is not None:
                fill = _zeg_fill_hex(avg_zeg)
                if fill:
                    _set_cell_shading(row.cells[ci], fill)
            _set_cell_text(row.cells[ci], txt)


def _fill_calc_detail(table, perf_rows: list[dict], bg_pct: str) -> None:
    _clear_rows(table, keep=1)
    for pr in perf_rows:
        z = pr["zeg"]
        kurs_days = (pr["kurs_h"] or 0) / 8.4
        mkt_days = (pr["marketing"] or 0) / 8.4
        lauf_days = (pr["lauf"] or 0) / 8.4
        row = table.add_row().cells
        _set_cell_text(row[0], MONTH_SHORT[pr["month"]])
        _set_cell_text(row[1], bg_pct)
        _set_cell_text(row[2], f"{pr['soll']:.1f}")
        _set_cell_text(row[3], _dash_val(pr["ferien"]))
        _set_cell_text(row[4], _dash_val(round(kurs_days, 2)) if kurs_days else "—")
        _set_cell_text(row[5], _dash_val(round(mkt_days, 2)) if mkt_days else "—")
        _set_cell_text(row[6], _dash_val(round(lauf_days, 2)) if lauf_days else "—")
        _set_cell_text(row[7], _dash_val(z.get("mgmt_t", 0)) if z.get("mgmt_t") else "—")
        _set_cell_text(row[8], _dash_val(pr["krank"]))
        _set_cell_text(row[9], f"= {z.get('prod_b', 0):g} T")
        _set_cell_text(row[10], _fmt_chf(pr["umsatz"]))
    formula = (
        "Formel: Soll-Tage (BG% × Arbeitstage Monat) − Abzüge = produktive Tage (B). "
        f"ZEG-B = Umsatz ÷ (ProdTage(B) × CHF {int(ZIEL)}). "
        "Krank-Tage werden aus ZEG-B herausgerechnet (Spalte C), nicht aus ZEG-B."
    )
    form_row = table.add_row()
    if len(form_row.cells) > 1:
        _merge_row_cells(form_row, 0, len(form_row.cells) - 1)
    _set_cell_text(form_row.cells[0], formula)
    for para in form_row.cells[0].paragraphs:
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.size = Pt(8)


def fill_hj1_template(
    ma: MAStammdaten,
    year: int,
    through_month: int,
    umsatz_all: dict,
    inputs_all: dict,
    bilat: BilatData | None,
    db: Session,
    dest_path: str,
) -> str:
    template_path = resolve_bilat_template(ma.name)
    shutil.copy(template_path, dest_path)
    doc = Document(dest_path)
    tables = _find_tables(doc)

    bg_pct = f"{round((ma.bg_pct or 0) * 100):.0f}%"
    period = _half_label(through_month, year)
    perf_range = _period_range(through_month, year)
    half_num = 1 if through_month <= 6 else 2
    ma_name = ma.display_name or ma.name

    perf_rows = _collect_performance(ma, year, through_month, umsatz_all, inputs_all, db)
    zeg_values = [r["zeg"]["zeg_b"] for r in perf_rows if r["zeg"].get("zeg_b") is not None]
    avg_zeg = sum(zeg_values) / len(zeg_values) if zeg_values else None

    fk = _find_fk(db, ma.team, ma.fk_username)
    fk_name = fk.full_name if fk and fk.full_name else None
    standort_funktion = _standort_funktion(ma)

    # Kopfdaten aus App (Stammdaten + FK)
    if "header" in tables:
        h = tables["header"]
        _set_cell_text(h.rows[1].cells[0], ma_name)
        _set_cell_text(h.rows[1].cells[1], standort_funktion)
        _set_cell_text(h.rows[1].cells[2], bg_pct)
        if len(h.rows[1].cells) > 4:
            _set_cell_text(h.rows[1].cells[4], period)
        if fk_name and len(h.rows) > 3:
            _set_cell_text(h.rows[3].cells[0], fk_name)

    if "perf_title" in tables:
        _set_cell_text(
            tables["perf_title"].rows[0].cells[0],
            f"  PERFORMANCE  {perf_range}  (ZEG-B: % des Ziels  |  Ziel = 100%)",
            section_header=True,
        )

    if "zeg_matrix" in tables:
        _fill_zeg_matrix(tables["zeg_matrix"], perf_rows, avg_zeg, bg_pct)

    if "qual_title" in tables:
        _set_cell_text(
            tables["qual_title"].rows[0].cells[0],
            f"  QUALITATIVE ZIELE {half_num}. HALBJAHR {year}",
            section_header=True,
        )

    qual_goal_names: list[str] = []
    if "qual_goals" in tables:
        qual_goal_names = _reset_qual_goal_values(tables["qual_goals"])

    if bilat and "ratings" in tables:
        t7 = tables["ratings"]
        kat_rows = {2: "a", 4: "b", 6: "c", 8: "d"}
        for row_idx, key in kat_rows.items():
            if row_idx >= len(t7.rows):
                continue
            self_v = getattr(bilat, f"kat_{key}_self", None)
            fk_v = getattr(bilat, f"kat_{key}_fk", None)
            comment = getattr(bilat, f"kat_{key}_comment", None) or ""
            if len(t7.rows[row_idx].cells) > 3:
                _set_rating_cell(t7.rows[row_idx].cells[2], self_v)
                _set_rating_cell(t7.rows[row_idx].cells[3], fk_v)
            if row_idx + 1 < len(t7.rows) and len(t7.rows[row_idx + 1].cells) > 1:
                _set_cell_text(t7.rows[row_idx + 1].cells[1], comment)
            _set_row_cant_split(t7.rows[row_idx], keep_with_next=True)
            if row_idx + 1 < len(t7.rows):
                _set_row_cant_split(t7.rows[row_idx + 1])

    if bilat and bilat.themen_ma and "themen" in tables and len(tables["themen"].rows) > 1:
        _set_cell_text(tables["themen"].rows[1].cells[0], bilat.themen_ma)

    elif "themen" in tables and len(tables["themen"].rows) > 1:
        _set_cell_text(tables["themen"].rows[1].cells[0], "")

    if bilat and bilat.vereinbarungen and "vereinbarungen" in tables and len(tables["vereinbarungen"].rows) > 1:
        _set_cell_text(tables["vereinbarungen"].rows[1].cells[0], bilat.vereinbarungen)
    elif "vereinbarungen" in tables and len(tables["vereinbarungen"].rows) > 1:
        _set_cell_text(tables["vereinbarungen"].rows[1].cells[0], "")

    if bilat and "abschluss" in tables:
        cells = tables["abschluss"].rows[0].cells
        if bilat.naechstes_bilat:
            _set_cell_text(cells[0], f"Nächstes Bilat-Datum: {bilat.naechstes_bilat}")
        if bilat.gespraechseindruck and len(cells) > 2:
            _set_cell_text(cells[2], f"Gesprächseindruck: {bilat.gespraechseindruck}")

    if "faktenblatt" in tables:
        _set_cell_text(
            tables["faktenblatt"].rows[0].cells[0],
            f"FAKTENBLATT  |  FK-intern  |  Nicht für Mitarbeiter/in | {ma_name}  |  {ma.team or '—'}  |  BG {bg_pct}  |  {period}",
            section_header=True,
        )

    if "perf_detail_title" in tables:
        _set_cell_text(
            tables["perf_detail_title"].rows[0].cells[0],
            f"  PERFORMANCE-DETAIL  {perf_range}",
            section_header=True,
        )

    if "perf_detail" in tables:
        _fill_perf_detail(tables["perf_detail"], perf_rows, year, perf_range, avg_zeg)

    if "calc_detail" in tables:
        _fill_calc_detail(tables["calc_detail"], perf_rows, bg_pct)

    if "leitfaden" in tables:
        t20 = tables["leitfaden"]
        points = _build_leitfaden_points(perf_range, qual_goal_names, bilat)
        _set_cell_text(t20.rows[0].cells[0], f"Performance-Einschätzung:\n{_performance_comment(avg_zeg, perf_range)}")
        if len(t20.rows) > 1:
            _set_cell_text(t20.rows[1].cells[0],
                           "Vorbereitete Gesprächspunkte:\n" + "\n".join(f"| {p}" for p in points))
        if len(t20.rows) > 2:
            _set_cell_text(t20.rows[2].cells[0],
                           "Vereinbarungen vorbereiten:\nKonkrete nächste Schritte für jedes offene / kritische Ziel festlegen.")

    doc.save(dest_path)
    return dest_path
