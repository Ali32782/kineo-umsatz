"""Tests für Bilateral-Export mit pro-MA-Vorlagen."""
from pathlib import Path

from bilat_hj1_export import fill_hj1_template, _find_tables, _zeg_fill_hex
from bilat_template_map import MA_BILAT_TEMPLATE, resolve_bilat_template, TEMPLATES_DIR
from docx import Document
from docx.oxml.ns import qn


def _cell_fill(cell) -> str | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _cell_xml_texts(cell) -> list[str]:
    import re
    return re.findall(r"<w:t[^>]*>([^<]*)</w:t>", cell._tc.xml)


def test_all_mapped_templates_exist():
    for ma_name, filename in MA_BILAT_TEMPLATE.items():
        path = TEMPLATES_DIR / filename
        assert path.is_file(), f"Vorlage fehlt für {ma_name}: {filename}"


def test_resolve_bilat_template_barbara():
    path = resolve_bilat_template("Barbara.V")
    assert path.name == "Bilat_Barbara_V_HJ1_2026.docx"


def test_template_has_expected_sections():
    doc = Document(resolve_bilat_template("Barbara.V"))
    tables = _find_tables(doc)
    assert "header" in tables
    assert "zeg_matrix" in tables
    assert "qual_goals" in tables
    assert "ratings" in tables
    assert "perf_detail" in tables
    assert "calc_detail" in tables
    assert "leitfaden" in tables


def test_generate_bilat_barbara(tmp_path):
    import os
    os.environ.setdefault("DATABASE_URL", "sqlite:///:memory:")
    from database import init_db, SessionLocal, MAStammdaten, UmsatzData

    init_db()
    db = SessionLocal()
    try:
        ma = db.query(MAStammdaten).filter_by(name="Barbara.V").first()
        if not ma:
            return
        db.add(UmsatzData(ma_name="Barbara.V", year=2026, month=3, umsatz=20000))
        db.commit()
        out = tmp_path / "out.docx"
        fill_hj1_template(ma, 2026, 6, {("Barbara.V", 3): 20000}, {}, None, db, str(out))
        assert out.is_file()
        doc = Document(out)
        tables = _find_tables(doc)
        assert "qual_goals" in tables
        goals = [r.cells[0].text for r in tables["qual_goals"].rows[1:3]]
        assert any("Schröpfen" in g or "Dokumentation" in g for g in goals)
    finally:
        db.close()


def test_read_qual_goals_and_extended_ratings():
    from bilat_hj1_export import _read_qual_goals_from_template, _read_rating_categories_from_template
    goals = _read_qual_goals_from_template("Noah.S")
    assert any("Movement Control" in g["name"] for g in goals)
    hanna = _read_rating_categories_from_template("Hanna.R")
    keys = {c["key"] for c in hanna}
    assert keys >= {"a", "b", "c", "d", "e", "f"}



def test_export_preserves_zeg_colors_and_ratings(tmp_path):
    import os
    os.environ.setdefault("DATABASE_URL", "sqlite:///:memory:")
    from database import init_db, SessionLocal, MAStammdaten, UmsatzData, BilatData

    init_db()
    db = SessionLocal()
    try:
        ma = db.query(MAStammdaten).filter_by(name="Hanna.R").first()
        if not ma:
            return
        for m in range(1, 7):
            db.add(UmsatzData(ma_name="Hanna.R", year=2026, month=m, umsatz=18000 + m * 500))
        bilat = BilatData(
            ma_name="Hanna.R", year=2026, period_label="1. HJ 2026",
            kat_a_self=4, kat_a_fk=3,
            kat_b_self=3, kat_b_fk=4,
            kat_c_self=5, kat_c_fk=3,
            kat_d_self=3, kat_d_fk=3,
        )
        db.add(bilat)
        db.commit()
        out = tmp_path / "hanna.docx"
        fill_hj1_template(ma, 2026, 6, {( "Hanna.R", m): 18000 + m * 500 for m in range(1, 7)}, {}, bilat, db, str(out))
        doc = Document(out)
        tables = _find_tables(doc)
        zeg_row = tables["zeg_matrix"].rows[1]
        assert _cell_fill(zeg_row.cells[1]) in ("FDEBD0", "D5F5E3", "FADBD8")
        assert _cell_fill(tables["zeg_matrix"].rows[0].cells[1]) == "2C3E50"

        ratings = tables["ratings"]
        ma_cell = ratings.rows[2].cells[2]
        assert "☑" in _cell_xml_texts(ma_cell)
        perf = tables["perf_detail"].rows[1].cells[1]
        assert _cell_fill(perf) in ("FDEBD0", "D5F5E3", "FADBD8")
    finally:
        db.close()
